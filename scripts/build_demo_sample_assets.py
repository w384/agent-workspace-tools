"""Build the first fictional PDF/DOCX sample set for the public-drive demo."""

from __future__ import annotations

import base64
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


PROJECT_ROOT = Path(__file__).resolve().parents[1]
SOURCE_ROOT = (
    PROJECT_ROOT / "work" / "demo" / "public-drive-ai-organizing" / "source"
)

PDF_DOCUMENTS = {
    "客户资料/星河食品品牌资料.pdf": (
        "星河食品有限公司品牌资料",
        (
            "星河食品有限公司定位为面向城市家庭的健康食品品牌。",
            "2026 春季新品传播统一使用“轻享春味”主题，主色为暖橙与清新绿。",
            "所有对外物料须保留品牌名称、产品名称和版本日期。",
        ),
    ),
    "报价合同/星河食品春季新品项目报价单.pdf": (
        "星河食品 2026 春季新品项目报价单",
        (
            "本报价覆盖策略、KV 设计、视频脚本和结项交付支持。",
            "交付文件应采用统一命名：客户-项目-文件类型-版本。",
            "项目归档时，报价与验收材料须保留在可追溯目录。",
        ),
    ),
    "输出稿/星河食品春季新品主视觉KV.pdf": (
        "星河食品春季新品主视觉 KV",
        (
            "主视觉交付版本：v1。",
            "文件用于 2026 春季新品上市整合传播项目的线上投放与线下陈列。",
            "验收前应确认品牌标识、产品名称、发布日期和版权素材来源。",
        ),
    ),
    "验收交付/2026春季新品项目验收清单.pdf": (
        "2026 春季新品项目验收清单",
        (
            "验收要求一：输出稿必须使用统一命名并存入验收交付目录。",
            "验收要求二：主视觉 KV、视频脚本、报价单和修改反馈应可追溯到同一项目。",
            "验收要求三：版权授权证明应由项目负责人单独核对。",
        ),
    ),
}

DOCX_DOCUMENTS = {
    "项目策划/2026春季新品整合传播方案.docx": (
        "2026 春季新品整合传播方案",
        (
            "项目目标：帮助星河食品有限公司完成春季新品上市整合传播。",
            "项目阶段：客户资料梳理、创意策划、制作输出、修改反馈、验收交付。",
            "文件整理要求：所有交付稿按客户、项目、文件类型和版本命名。",
        ),
    ),
    "输出稿/星河食品春季新品视频脚本.docx": (
        "星河食品春季新品视频脚本",
        (
            "第一段：春日餐桌场景展示新品轻享体验。",
            "第二段：产品卖点和品牌信息同步出现。",
            "第三段：交付版本应标记为 v1，并在验收前保留修改记录。",
        ),
    ),
    "修改反馈/客户第三轮修改意见.docx": (
        "客户第三轮修改意见",
        (
            "请将主视觉标题缩短，并保持产品名称清晰可读。",
            "视频脚本结尾增加新品上市日期提示。",
            "修改完成后，请更新输出稿并保留反馈关联。",
        ),
    ),
    "版权授权证明/内部法务评审意见.docx": (
        "内部法务评审意见",
        (
            "仅限项目负责人查阅：库存图片授权期限需要覆盖春季投放周期。",
            "若素材授权范围发生变化，应在交付前重新核对并更新证明。",
            "本文件不对项目执行成员开放查询权限。",
        ),
    ),
}

ONE_PIXEL_JPEG = base64.b64decode(
    "/9j/4AAQSkZJRgABAQAAAQABAAD/2wBDAP//////////////////////////////////////////////////////////////////////////////////////"
    "2wBDAf//////////////////////////////////////////////////////////////////////////////////////wAARCAABAAEDASIAAhEBAxEB/8QAFQABAQAAAAAAAAAAAAAAAAAAAAX/xAAUEAEAAAAAAAAAAAAAAAAAAAAA/9oADAMBAAIQAxAAAAH/xAAUEAEAAAAAAAAAAAAAAAAAAAAA/9oACAEBAAEFAqf/xAAUEQEAAAAAAAAAAAAAAAAAAAAA/9oACAEDAQE/Aaf/xAAUEQEAAAAAAAAAAAAAAAAAAAAA/9oACAECAQE/Aaf/xAAUEAEAAAAAAAAAAAAAAAAAAAAA/9oACAEBAAY/Ap//xAAUEAEAAAAAAAAAAAAAAAAAAAAA/9oACAEBAAE/IX//2gAMAwEAAgADAAAAEP/EABQRAQAAAAAAAAAAAAAAAAAAABD/2gAIAQMBAT8QH//EABQRAQAAAAAAAAAAAAAAAAAAABD/2gAIAQIBAT8QH//EABQQAQAAAAAAAAAAAAAAAAAAABD/2gAIAQEAAT8QH//Z"
)


def main() -> None:
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    for relative_path, (title, paragraphs) in PDF_DOCUMENTS.items():
        _build_pdf(SOURCE_ROOT / relative_path, title, paragraphs)
    for relative_path, (title, paragraphs) in DOCX_DOCUMENTS.items():
        _build_docx(SOURCE_ROOT / relative_path, title, paragraphs)

    image_path = SOURCE_ROOT / "原始素材/星河食品新品KV参考.jpg"
    image_path.parent.mkdir(parents=True, exist_ok=True)
    image_path.write_bytes(ONE_PIXEL_JPEG)

    project_path = SOURCE_ROOT / "工程文件/春季新品主视觉工程.aep"
    project_path.parent.mkdir(parents=True, exist_ok=True)
    project_path.write_text(
        "Fictional After Effects placeholder. Metadata-only demo asset; not parsed.\n",
        encoding="utf-8",
    )


def _build_pdf(path: Path, title: str, paragraphs: tuple[str, ...]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    styles = getSampleStyleSheet()
    title_style = styles["Title"].clone("DemoChineseTitle")
    title_style.fontName = "STSong-Light"
    title_style.fontSize = 18
    body_style = styles["BodyText"].clone("DemoChineseBody")
    body_style.fontName = "STSong-Light"
    body_style.fontSize = 11
    body_style.leading = 18
    document = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        leftMargin=24 * mm,
        rightMargin=24 * mm,
        topMargin=24 * mm,
        bottomMargin=24 * mm,
        title=title,
        author="公共盘 AI 整理 DEMO",
    )
    story = [Paragraph(title, title_style), Spacer(1, 8 * mm)]
    story.extend(Paragraph(paragraph, body_style) for paragraph in paragraphs)
    document.build(story)


def _build_docx(path: Path, title: str, paragraphs: tuple[str, ...]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    section = document.sections[0]
    section.top_margin = Pt(56)
    section.bottom_margin = Pt(56)
    document.core_properties.title = title
    document.core_properties.author = "公共盘 AI 整理 DEMO"
    heading = document.add_heading(title, level=1)
    _apply_chinese_font(heading, 18)
    for paragraph_text in paragraphs:
        paragraph = document.add_paragraph(paragraph_text)
        _apply_chinese_font(paragraph, 11)
    document.save(path)


def _apply_chinese_font(paragraph, point_size: int) -> None:
    for run in paragraph.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(point_size)


if __name__ == "__main__":
    main()
