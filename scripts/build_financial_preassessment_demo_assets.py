"""Build fictional financial-preassessment demo assets and rule fixtures."""

from __future__ import annotations

import hashlib
import json
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


PROJECT_ROOT = Path(__file__).resolve().parents[1]
DEMO_ROOT = PROJECT_ROOT / "work" / "demo" / "financial-preassessment"
SOURCE_ROOT = DEMO_ROOT / "source"
RULES_PATH = DEMO_ROOT / "rules" / "demo-bank-rules-v1.json"
IMPORT_MANIFEST_PATH = DEMO_ROOT / "import-manifest.json"
DISCLAIMER = "仅供资料完整度与规则匹配演示参考，不参与贷款申请、审批、授信、额度测算或金融产品销售。"

PDF_DOCUMENTS = {
    "客户模拟资料/收入情况说明.pdf": (
        "收入情况说明",
        (
            "一、企业概况",
            "深圳市海川智能装备制造有限公司（统一社会信用代码 91440300MA5F2K9X4Q），成立于2019年3月，注册资本人民币2,000万元（实缴1,500万元），注册地址位于深圳市宝安区西乡街道工业区12栋。公司主营工业自动化装配检测线、工业机器人上下料单元及非标自动化设备的设计、制造与销售，并配套提供产线工艺调试技术服务，现有员工86人。",
            "二、近年经营收入情况",
            "2022年度营业收入2,860万元；2023年度营业收入4,118万元，同比增长44.0%；2024年度营业收入4,860万元，同比增长18.0%。收入构成以设备销售收入为主（约88%），技术服务收入为辅（约12%）。",
            "三、开票与纳税申报口径",
            "2024年度开具增值税发票（销项）不含税金额4,580万元；增值税纳税申报销售额4,720万元。两者差异140万元，主要为2024年11月交付的研发服务合同尚未开具发票，预计于次年1月补开。2024年度取得进项发票不含税金额2,940万元。",
            "四、收入真实性说明",
            "上述收入与主要客户销售合同、账户回款记录可相互印证；主要客户回款与确认收入时间偏差在信用期内，回款情况正常，未出现异常大额冲减。",
            "五、声明",
            "本企业确认上述收入及申报数据真实、准确、完整，相关原始凭证（销售合同、发票、发货单、验收单）留档可查。",
            DISCLAIMER,
        ),
    ),
    "客户模拟资料/资金流摘要.pdf": (
        "资金流摘要",
        (
            "一、说明",
            "本摘要基于企业主要结算账户（基本存款账户及一般结算账户）2024年1月至12月的实际收付流水整理，供经营与资金情况核实参考。",
            "二、经营性资金流入",
            "2024年度经营性资金流入合计5,120万元，其中：销售回款4,760万元，其他经营性流入（退货款、押金退回、政府补贴等）360万元。销售回款率（回款/营业收入）为97.9%，回款总体及时。",
            "三、经营性资金流出",
            "2024年度经营性资金流出合计4,440万元，其中：原材料及设备采购支出3,120万元，职工薪酬780万元，税费缴纳320万元，其他经营支出220万元。",
            "四、经营性净现金流",
            "2024年度经营性净现金流约680万元，经营自身造血能力稳定，未依赖融资性流入维持日常经营。",
            "五、大额资金往来说明",
            "2024年12月收到客户A公司年终结算款项480万元，对应2024年11月交付验收的自动化装配线合同尾款，与销售合同及验收单一致，属正常经营性回款。",
            "六、资金流与财务数据匹配",
            "销售回款与营业收入、应收账款变动总体匹配：2024年末应收账款1,480万元，其中账龄一年以内占比约90%。",
            DISCLAIMER,
        ),
    ),
    "客户模拟资料/补充材料清单.pdf": (
        "补充材料清单",
        (
            "一、已提供材料",
            "1. 资料概览与授权说明；2. 收入情况说明；3. 资金流摘要；4. 资产负债说明；5. 经营情况说明；6. 营业执照及法定代表人身份证明（另附）。",
            "二、待补充材料",
            "1. 连续两年经审计的经营证明（近两年审计报告或年度纳税申报汇总）；2. 可核验的外部资质说明（如高新技术企业证书、行业认证等）；3. 主要客户及供应商名录（含合作年限）。",
            "三、特殊事项补充说明",
            "1. 2024年12月大额一次性回款480万元，为与大客户A公司的年度结算款，凭证齐全；2. 2024年度新签合同额约5,600万元，其中在手未交付合同约3,200万元，相关合同清单可提供备查。",
            "四、联系人",
            "资料提供与核验联系人：财务部，联系电话以授权说明所载为准。",
            DISCLAIMER,
        ),
    ),
    "规则依据/演示规则适用说明.pdf": (
        "演示规则适用说明",
        (
            "此文件仅说明演示规则的使用边界；规则正文以 demo-bank-rules-v1.json 的版本指纹为准。",
            "所有银行名称、规则条件和结果等级均为虚构示例，不代表任何真实银行或金融机构要求。",
            DISCLAIMER,
        ),
    ),
}

DOCX_DOCUMENTS = {
    "客户模拟资料/资料概览与授权说明.docx": (
        "资料概览与授权说明",
        (
            "一、企业基本信息",
            "企业名称：深圳市海川智能装备制造有限公司；统一社会信用代码：91440300MA5F2K9X4Q；成立时间：2019年3月；注册资本：人民币2,000万元（实缴1,500万元）；注册地址：深圳市宝安区西乡街道工业区12栋；法定代表人：林海峰；所属行业：通用设备制造业（工业自动化装备）；员工人数：86人。",
            "二、主营业务",
            "公司面向电子制造、汽车零部件、家电、新能源及食品包装行业，提供自动化装配检测线、工业机器人上下料单元、非标自动化设备的设计、制造与安装调试，并提供产线工艺调试与设备维保服务。",
            "三、本次提交资料清单",
            "1. 收入情况说明；2. 资金流摘要；3. 资产负债说明；4. 经营情况说明；5. 补充材料清单；6. 营业执照及法定代表人身份证明。",
            "四、授权说明",
            "本企业授权将上述资料提交至企业资料预评估平台，用于资料完整度、有效性及规则匹配情况的演示评估；本次评估结果仅反映资料与示例规则的匹配程度，不构成贷款申请、授信审批、额度测算或金融产品销售的授权依据。如后续开展真实业务，将另行提供经盖章确认的正式申请材料。",
            "五、联系人",
            "财务负责人：张敏；联系方式：以企业登记信息所载为准。",
            DISCLAIMER,
        ),
    ),
    "客户模拟资料/资产负债说明.docx": (
        "资产负债说明",
        (
            "一、资产负债表摘要（截至2024年12月31日）",
            "资产总计7,200万元：其中货币资金860万元、应收账款1,480万元、存货1,120万元（原材料480万元、在产品260万元、产成品380万元）、固定资产净值2,340万元、预付款项及其他资产1,400万元。",
            "负债及所有者权益总计7,200万元：其中短期借款1,200万元（银行抵押贷款800万元、信用贷款400万元）、应付账款980万元、其他流动负债（合同负债、应付职工薪酬、应交税费等）1,220万元、所有者权益（净资产）3,800万元。",
            "二、偿债能力指标",
            "资产负债率约47.2%；流动比率约1.7；短期借款均有对应经营周转安排，未出现逾期或展期情况。",
            "三、应收账款账龄",
            "应收账款1,480万元中：账龄一年以内约1,332万元（占比90%）、一至二年约118万元（占比8%）、二年以上约30万元（占比2%），历史坏账率低。",
            "四、存货构成",
            "存货以在制项目和标准备件为主，产成品对应已签合同订单，不存在明显积压；计提存货跌价准备符合企业会计政策。",
            "五、担保与或有事项",
            "企业无对外担保，无未决重大诉讼，无资产被查封、冻结或质押（短期借款抵押资产除外），无重大未披露或有负债。",
            DISCLAIMER,
        ),
    ),
    "客户模拟资料/经营情况说明.docx": (
        "经营情况说明",
        (
            "一、主营业务",
            "公司为工业自动化装备制造商，主要产品包括自动化装配检测线、工业机器人上下料单元及非标自动化设备，并提供配套工艺调试服务。产品主要用于电子制造、汽车零部件、家电、新能源及食品包装行业的产线自动化升级。",
            "二、经营连续性",
            "公司自2019年成立以来持续经营，营业收入逐年增长：2022年2,860万元、2023年4,118万元、2024年4,860万元，连续正常纳税，经营记录完整。",
            "三、主要客户",
            "前五大客户收入占比约62%：客户A（电子代工）18%、客户B（汽车零部件）15%、客户C（家电）12%、客户D（新能源）9%、客户E（食品包装）8%；与主要客户合作年限多在三年以上，合作关系稳定。",
            "四、主要供应商",
            "主要供应商8家，前三大供应商采购占比约45%，涉及标准件、钣金加工件及电气元件，供应渠道稳定。",
            "五、成长性与在手订单",
            "2024年度新签合同额约5,600万元；截至2024年末在手未交付合同约3,200万元，2025年一季度新增订单约1,100万元。",
            "六、资金需求说明",
            "本次拟申请流动资金用于扩大原材料备货及满足在手订单交付的资金周转需求，资金用途明确，还款来源为经营回款。",
            DISCLAIMER,
        ),
    ),
    "敏感资料/内部资料核验说明.docx": (
        "内部资料核验说明（模拟敏感资料）",
        (
            "本文件是虚构的内部核验说明，仅用于演示权限负向控制。",
            "项目普通成员不应查询本文件；只有被显式授予权限的负责人可引用。",
            DISCLAIMER,
        ),
    ),
}


def main() -> None:
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    for relative_path, (title, paragraphs) in PDF_DOCUMENTS.items():
        _build_pdf(SOURCE_ROOT / relative_path, title, paragraphs)
    for relative_path, (title, paragraphs) in DOCX_DOCUMENTS.items():
        _build_docx(SOURCE_ROOT / relative_path, title, paragraphs)
    _write_rule_fixture()
    _write_import_manifest()


def _write_rule_fixture() -> None:
    payload = {
        "source_type": "demo_fixture",
        "scenario": "finance_profile_matching",
        "rule_set_name": "演示银行规则样例",
        "assessment_rule_id": "demo-bank-a-complete",
        "version_label": "demo-2026-08-14",
        "disclaimer": DISCLAIMER,
        "rules": [
            {
                "rule_id": "demo-bank-a-complete",
                "bank_label": "示例银行A",
                "result_level": "MATCH",
                "requirements": [
                    {
                        "rule_id": "demo-bank-a-income",
                        "material_key": "income_statement",
                        "label": "收入情况说明",
                    },
                    {
                        "rule_id": "demo-bank-a-cashflow",
                        "material_key": "cashflow_summary",
                        "label": "资金流摘要",
                    },
                    {
                        "rule_id": "demo-bank-a-assets",
                        "material_key": "asset_liability_statement",
                        "label": "资产负债说明",
                    },
                ],
            },
            {
                "rule_id": "demo-bank-b-supplement",
                "bank_label": "示例银行B",
                "result_level": "POSSIBLE",
                "requirements": [
                    {
                        "rule_id": "demo-bank-b-income",
                        "material_key": "income_statement",
                        "label": "收入情况说明",
                    },
                    {
                        "rule_id": "demo-bank-b-business",
                        "material_key": "business_profile",
                        "label": "经营情况说明",
                    },
                    {
                        "rule_id": "demo-bank-b-supplement",
                        "material_key": "supplement_material_list",
                        "label": "补充材料清单",
                    },
                ],
            },
            {
                "rule_id": "demo-bank-c-history",
                "bank_label": "示例银行C",
                "result_level": "NOT_MATCH",
                "requirements": [
                    {
                        "rule_id": "demo-bank-c-history",
                        "material_key": "two_year_business_history",
                        "label": "连续两年经营证明",
                    },
                    {
                        "rule_id": "demo-bank-c-license",
                        "material_key": "external_qualification",
                        "label": "可核验的外部资质说明",
                    },
                ],
            },
            {
                "rule_id": "demo-bank-d-micro",
                "bank_label": "示例银行D",
                "result_level": "POSSIBLE",
                "requirements": [
                    {
                        "rule_id": "demo-bank-d-profile",
                        "material_key": "customer_profile",
                        "label": "资料概览与授权说明",
                    },
                    {
                        "rule_id": "demo-bank-d-income",
                        "material_key": "income_statement",
                        "label": "收入情况说明",
                    },
                    {
                        "rule_id": "demo-bank-d-business",
                        "material_key": "business_profile",
                        "label": "经营情况说明",
                    },
                    {
                        "rule_id": "demo-bank-d-supplement",
                        "material_key": "supplement_material_list",
                        "label": "补充材料清单",
                    },
                ],
            },
            {
                "rule_id": "demo-bank-e-cashflow",
                "bank_label": "示例银行E",
                "result_level": "POSSIBLE",
                "requirements": [
                    {
                        "rule_id": "demo-bank-e-cashflow",
                        "material_key": "cashflow_summary",
                        "label": "资金流摘要",
                    },
                    {
                        "rule_id": "demo-bank-e-assets",
                        "material_key": "asset_liability_statement",
                        "label": "资产负债说明",
                    },
                    {
                        "rule_id": "demo-bank-e-income",
                        "material_key": "income_statement",
                        "label": "收入情况说明",
                    },
                ],
            },
        ],
    }
    canonical = json.dumps(payload, ensure_ascii=False, sort_keys=True, separators=(",", ":"))
    payload["content_fingerprint"] = "sha256:" + hashlib.sha256(
        canonical.encode("utf-8")
    ).hexdigest()
    RULES_PATH.parent.mkdir(parents=True, exist_ok=True)
    RULES_PATH.write_text(
        json.dumps(payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8"
    )


def _write_import_manifest() -> None:
    payload = {
        "source_type": "demo_fixture",
        "scenario": "finance_profile_matching",
        "assets": [
            {
                "relative_path": "客户模拟资料/资料概览与授权说明.docx",
                "material_key": "customer_profile",
            },
            {
                "relative_path": "客户模拟资料/收入情况说明.pdf",
                "material_key": "income_statement",
            },
            {
                "relative_path": "客户模拟资料/资金流摘要.pdf",
                "material_key": "cashflow_summary",
            },
            {
                "relative_path": "客户模拟资料/资产负债说明.docx",
                "material_key": "asset_liability_statement",
            },
            {
                "relative_path": "客户模拟资料/经营情况说明.docx",
                "material_key": "business_profile",
            },
            {
                "relative_path": "客户模拟资料/补充材料清单.pdf",
                "material_key": "supplement_material_list",
            },
        ],
    }
    IMPORT_MANIFEST_PATH.write_text(
        json.dumps(payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8"
    )


def _build_pdf(path: Path, title: str, paragraphs: tuple[str, ...]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    styles = getSampleStyleSheet()
    title_style = styles["Title"].clone("FinancialDemoTitle")
    title_style.fontName = "STSong-Light"
    title_style.fontSize = 18
    body_style = styles["BodyText"].clone("FinancialDemoBody")
    body_style.fontName = "STSong-Light"
    body_style.fontSize = 11
    body_style.leading = 18
    document = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        leftMargin=24 * mm,
        rightMargin=24 * mm,
        topMargin=24 * mm,
        bottomMargin=24 * mm,
        title=title,
        author="企业资料资产化与场景知识库 DEMO",
    )
    story = [Paragraph(title, title_style), Spacer(1, 8 * mm)]
    story.extend(Paragraph(paragraph, body_style) for paragraph in paragraphs)
    document.build(story)
    _pin_pdf_datetime(path)


def _pin_pdf_datetime(path: Path) -> None:
    """Freeze CreationDate/ModDate so rebuilds do not churn git history.

    reportlab stamps the current time into the PDF info dict, so every rebuild
    changes bytes even with identical content. The replacement is same-length
    (both formats are D:YYYYMMDDHHMMSS+08'00') so xref byte offsets stay valid.
    """
    raw = path.read_bytes()
    fixed = b"D:20260814000000+08'00'"
    raw = re.sub(rb"/CreationDate \(D:[^)]*\)", b"/CreationDate (" + fixed + b")", raw)
    raw = re.sub(rb"/ModDate \(D:[^)]*\)", b"/ModDate (" + fixed + b")", raw)
    path.write_bytes(raw)


def _build_docx(path: Path, title: str, paragraphs: tuple[str, ...]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.core_properties.title = title
    document.core_properties.author = "企业资料资产化与场景知识库 DEMO"
    fixed_date = datetime(2026, 8, 14, 10, 0, 13)
    document.core_properties.created = fixed_date
    document.core_properties.modified = fixed_date
    heading = document.add_heading(title, level=1)
    _apply_chinese_font(heading, 18)
    for paragraph_text in paragraphs:
        paragraph = document.add_paragraph(paragraph_text)
        _apply_chinese_font(paragraph, 11)
    document.save(path)


def _apply_chinese_font(paragraph, point_size: int) -> None:
    for run in paragraph.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(point_size)


if __name__ == "__main__":
    main()
